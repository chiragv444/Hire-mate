import os
import json
from typing import Dict, List, Any, Optional
from datetime import datetime
import PyPDF2
import docx
import pdfplumber
import re

# Try to import langchain dependencies, fallback to None if not available
try:
    from langchain.llms import OpenAI  # unused but retained to avoid breaking imports elsewhere
    from langchain.chat_models import ChatOpenAI
    from langchain.prompts import PromptTemplate
    from langchain.output_parsers import PydanticOutputParser
    from pydantic import BaseModel, Field
    LANGCHAIN_AVAILABLE = True
except ImportError:
    print("LangChain not available, using basic parsing only")
    LANGCHAIN_AVAILABLE = False
    # Create dummy classes for type hints
    class BaseModel:
        def dict(self):
            return {}
    class Field:
        def __init__(self, *args, **kwargs):
            pass

# ----------------------------
# Pydantic models (same output schema)
# ----------------------------

class PersonalInfo(BaseModel):
    name: Optional[str] = Field(description="Full name of the person")
    email: Optional[str] = Field(description="Email address")
    phone: Optional[str] = Field(description="Phone number")
    location: Optional[str] = Field(description="Location/Address")
    linkedin: Optional[str] = Field(description="LinkedIn profile URL")
    github: Optional[str] = Field(description="GitHub profile URL")
    website: Optional[str] = Field(description="Personal website URL")

class Education(BaseModel):
    degree: Optional[str] = Field(description="Degree name")
    institution: Optional[str] = Field(description="Institution name")
    graduation_year: Optional[str] = Field(description="Graduation year")
    gpa: Optional[str] = Field(description="GPA if mentioned")
    field_of_study: Optional[str] = Field(description="Field of study/Major")
    # Optional fields that many resumes include; harmless if unused by the rest of your API:
    location: Optional[str] = Field(description="Institution location")
    start_date: Optional[str] = Field(description="Start date")
    end_date: Optional[str] = Field(description="End date")

class Experience(BaseModel):
    title: Optional[str] = Field(description="Job title")
    company: Optional[str] = Field(description="Company name")
    location: Optional[str] = Field(description="Job location")
    start_date: Optional[str] = Field(description="Start date")
    end_date: Optional[str] = Field(description="End date or 'Present'")
    duration: Optional[str] = Field(description="Duration of employment")
    description: List[str] = Field(description="Job responsibilities and achievements", default_factory=list)

class Project(BaseModel):
    name: Optional[str] = Field(description="Project name")
    description: Optional[str] = Field(description="Project description")
    technologies: List[str] = Field(description="Technologies used", default_factory=list)
    url: Optional[str] = Field(description="Project URL if available")

class Certification(BaseModel):
    name: str = Field(description="Certification name")
    issuer: Optional[str] = Field(description="Issuing organization")
    date: Optional[str] = Field(description="Date obtained")
    expiry: Optional[str] = Field(description="Expiry date if applicable")

class ParsedResumeStructure(BaseModel):
    personal_info: PersonalInfo = Field(description="Personal information")
    summary: Optional[str] = Field(description="Professional summary or objective")
    skills: Dict[str, List[str]] = Field(description="Skills categorized by type (technical, soft, etc.)")
    experience: List[Experience] = Field(description="Work experience", default_factory=list)
    education: List[Education] = Field(description="Educational background", default_factory=list)
    projects: List[Project] = Field(description="Projects", default_factory=list)
    certifications: List[Certification] = Field(description="Certifications", default_factory=list)
    languages: List[str] = Field(description="Languages spoken", default_factory=list)
    awards: List[str] = Field(description="Awards and achievements", default_factory=list)

# ----------------------------
# Regex & parsing helpers
# ----------------------------

CID_RE = re.compile(r"\(cid:\d+\)")
MULTISPACE = re.compile(r"[ \t]+")
# Normalize many bullets to "- "
BULLETS = ["•", "◦", "▪", "■", "–", "—", "-", "●", "∙", "·", "►", "♦"]

URL_RE = re.compile(r"\b(?:https?://|www\.)[^\s)]+", re.IGNORECASE)
EMAIL_RE = re.compile(r"\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b")
PHONE_RE = re.compile(
    r"(?:\+?1[\s\-\.]?)?(?:\(?\d{3}\)?[\s\-\.]?)\d{3}[\s\-\.]?\d{4}"
)
LINKEDIN_ANY_RE = re.compile(r"(?:https?://|www\.)?linkedin\.com/[^\s|]+", re.IGNORECASE)
GITHUB_ANY_RE = re.compile(r"(?:https?://|www\.)?github\.com/[^\s|]+", re.IGNORECASE)

MONTHS = {
    'jan':1, 'january':1, 'feb':2, 'february':2, 'mar':3, 'march':3, 'apr':4, 'april':4,
    'may':5, 'jun':6, 'june':6, 'jul':7, 'july':7, 'aug':8, 'august':8, 'sep':9, 'sept':9,
    'september':9, 'oct':10, 'october':10, 'nov':11, 'november':11, 'dec':12, 'december':12
}

# US states & Canadian provinces (abbrev) for better location detection
US_CA_ABBR = {
    "AL","AK","AZ","AR","CA","CO","CT","DC","DE","FL","GA","HI","IA","ID","IL","IN","KS","KY","LA","MA","MD","ME","MI","MN","MO","MS","MT","NC","ND","NE","NH","NJ","NM","NV","NY","OH","OK","OR","PA","RI","SC","SD","TN","TX","UT","VA","VT","WA","WI","WV",
    "AB","BC","MB","NB","NL","NS","NT","NU","ON","PE","QC","SK","YT"
}

# Date spans: "Jan 2021 – Mar 2024", "2020 - 2022", "03/2020 to 06/2021", "2019 to Present"
DATE_SPAN_RE = re.compile(
    r"(?P<start>(?:"
    r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}"
    r"|"
    r"\d{1,2}/\d{4}"
    r"|"
    r"\d{4}"
    r"))\s*(?:[-–—]|to)\s*(?P<end>(?:"
    r"(?:Present|Now|Current|present|now|current)"
    r"|"
    r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}"
    r"|"
    r"\d{1,2}/\d{4}"
    r"|"
    r"\d{4}"
    r"))"
)

SECTION_ALIASES = {
    # summaries / profiles
    "summary": ["summary", "objective", "profile", "professional profile", "about", "highlights", "career summary", "professional summary"],
    # skills / competencies (broad, multi-domain)
    "skills": [
        "skills", "technical skills", "core skills", "competencies", "strengths",
        "technical competencies", "areas of expertise", "expertise", "capabilities",
        "key skills", "core competencies"
    ],
    # experience / employment
    "experience": [
        "experience", "work experience", "professional experience", "employment history",
        "work history", "relevant experience", "professional history", "career history",
        "industry experience", "volunteer experience", "volunteering"
    ],
    # education / training
    "education": ["education", "academic background", "academics", "education & training", "training", "coursework", "courses"],
    # projects / case studies / portfolios
    "projects": ["projects", "personal projects", "selected projects", "notable projects", "portfolio", "case studies", "engagements"],
    # certifications / licenses
    "certifications": ["certifications", "licenses", "licenses & certifications", "licences & certifications", "licences", "licenses and certifications"],
    # languages (spoken)
    "languages": ["languages", "language proficiency"],
    # awards / honors / publications
    "awards": ["awards", "honors", "honours", "achievements", "recognition", "publications", "press", "media"],
    # leadership / activities / affiliations
    "leadership": ["leadership", "leadership experience", "activities", "affiliations", "memberships"],
    # references
    "references": ["references"]
}

def _normalize_text_keep_layout(text: str) -> str:
    """Fix common PDF artifacts, keep newlines/bullets."""
    text = CID_RE.sub("", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    for b in BULLETS:
        text = text.replace(b, "- ")
    # collapse whitespace per line but keep line breaks
    lines = [MULTISPACE.sub(" ", ln).strip() for ln in text.split("\n")]
    return "\n".join(lines)

def _heading_key(line: str) -> Optional[str]:
    s = line.strip().lower().rstrip(":").strip()
    for key, names in SECTION_ALIASES.items():
        for n in names:
            if s == n:
                return key
    return None

def _split_sections(text: str) -> Dict[str, List[str]]:
    sections: Dict[str, List[str]] = {}
    current = "header"
    sections[current] = []
    for ln in text.split("\n"):
        k = _heading_key(ln)
        if k:
            current = k
            sections.setdefault(current, [])
            continue
        sections.setdefault(current, []).append(ln)
    return sections

def _parse_month_year(s: str) -> Optional[datetime]:
    if not s:
        return None
    s = s.strip()
    # "YYYY"
    if re.fullmatch(r"\d{4}", s):
        return datetime(int(s), 1, 1)
    # "MM/YYYY"
    m = re.fullmatch(r"(\d{1,2})/(\d{4})", s)
    if m:
        mm = max(1, min(12, int(m.group(1))))
        return datetime(int(m.group(2)), mm, 1)
    # "Mon YYYY"
    m = re.match(r"([A-Za-z]{3,9})\.?\s+(\d{4})", s)
    if m:
        month = MONTHS.get(m.group(1).lower())
        if month:
            return datetime(int(m.group(2)), month, 1)
    return None

def _months_between(start: Optional[datetime], end: Optional[datetime]) -> int:
    if not start:
        return 0
    end = end or datetime.utcnow()
    return max(0, (end.year - start.year) * 12 + (end.month - start.month))

def _collect_blocks(lines: List[str]) -> List[List[str]]:
    blocks = []
    cur: List[str] = []
    for ln in lines:
        if ln.strip() == "":
            if cur:
                blocks.append(cur)
                cur = []
            continue
        cur.append(ln)
    if cur:
        blocks.append(cur)
    return blocks

def _parse_bullets(lines: List[str]) -> List[str]:
    out = []
    for ln in lines:
        s = ln.strip()
        if not s or s in ("-", "–", "—"):
            continue
        if s.startswith("- "):
            s = s[2:].strip()
        out.append(s)
    return out

# ----------------------------
# Rule-based extraction
# ----------------------------

def _guess_name_from_header(header_lines: List[str]) -> Optional[str]:
    for ln in header_lines:
        s = ln.strip()
        if not s:
            continue
        if any(ch.isdigit() for ch in s):
            continue
        if 2 <= len(s.split()) <= 8 and 2 <= len(s) <= 80:
            return s
    return None

def _clean_url(u: str) -> str:
    if not u:
        return u
    # remove spaces and zero-widths inside URLs (common PDF issue)
    u = re.sub(r"\s+", "", u)
    u = u.replace("\u200b", "").replace("\u200c", "").replace("\u200d", "")
    # ensure scheme
    if u.startswith("www."):
        u = "https://" + u
    if not u.lower().startswith(("http://", "https://", "mailto:")):
        u = "https://" + u
    # strip trailing punctuation
    u = u.rstrip(").,;")
    return u

# broad location extraction: City, ST/Province [, Country]
CITY_STATE_RE = re.compile(
    r"\b([A-Za-z][A-Za-z .'\-]+),\s*([A-Z]{2})(?:,\s*(?:USA|United States|Canada))?\b"
)
CITY_PROV_COUNTRY_RE = re.compile(
    r"\b([A-Za-z][A-Za-z .'\-]+),\s*([A-Za-z][A-Za-z .'\-]+)(?:,\s*(Canada|USA|United States))?\b"
)

def _find_location_candidates(text: str) -> List[str]:
    cands: List[str] = []
    # 1) City, ST (Toronto, ON) / (New York, NY)
    for m in CITY_STATE_RE.finditer(text):
        city, abbr = m.group(1).strip(), m.group(2).strip().upper()
        if abbr in US_CA_ABBR:
            cands.append(f"{city}, {abbr}")
    # 2) City, Province/Country words (Toronto, Canada)
    for m in CITY_PROV_COUNTRY_RE.finditer(text):
        city, region, country = m.group(1).strip(), m.group(2).strip(), (m.group(3) or "").strip()
        if len(region.split()) <= 4 and len(city) <= 64:
            loc = f"{city}, {region}"
            if country and country not in loc:
                loc = f"{loc}, {country}"
            cands.append(loc)
    # de-duplicate preserving order
    seen = set()
    out = []
    for x in cands:
        key = x.lower()
        if key not in seen:
            seen.add(key)
            out.append(x)
    return out

def _majority_vote_location(header_lines: List[str], all_text: str, exp: List[Dict[str, Any]], edu: List[Dict[str, Any]]) -> Optional[str]:
    votes: List[str] = []
    votes += _find_location_candidates(" ".join(header_lines))  # header likely "current" location
    for e in exp or []:
        if e.get("location"):
            votes.append(e["location"])
    for ed in edu or []:
        if ed.get("location"):
            votes.append(ed["location"])
    first_lines = "\n".join(all_text.split("\n")[:30])
    votes += _find_location_candidates(first_lines)

    if not votes:
        return None
    norm = [v.strip() for v in votes if v and isinstance(v, str)]
    if not norm:
        return None
    counts = {}
    for v in norm:
        counts[v] = counts.get(v, 0) + 1
    norm_sorted = sorted(counts.items(), key=lambda kv: (-kv[1], norm.index(kv[0])))
    return norm_sorted[0][0] if norm_sorted else None

def _extract_personal_info(sections: Dict[str, List[str]]) -> Dict[str, Any]:
    header = sections.get("header", [])
    all_text = "\n".join(sum(sections.values(), []))

    email = EMAIL_RE.search(all_text)
    phone = PHONE_RE.search(all_text)

    # Try links anywhere in the doc (with cleaning)
    linkedin = None
    github = None
    website = None

    m = LINKEDIN_ANY_RE.search(all_text)
    if m:
        linkedin = _clean_url(m.group(0))

    m = GITHUB_ANY_RE.search(all_text)
    if m:
        github = _clean_url(m.group(0))

    urls = [_clean_url(u) for u in URL_RE.findall(all_text)]
    if urls:
        for u in urls:
            lu = u.lower()
            if "linkedin.com" in lu:
                linkedin = linkedin or u
                continue
            if "github.com" in lu:
                github = github or u
                continue
            if "mailto:" in lu:
                continue
            website = website or u

    # Location: inferred later with majority vote
    location = None
    name = _guess_name_from_header(header)

    return {
        "name": name,
        "email": email.group(0) if email else None,
        "phone": phone.group(0) if phone else None,
        "location": location,
        "linkedin": linkedin,
        "github": github,
        "website": website,
    }

def _extract_experience(sections: Dict[str, List[str]]) -> List[Dict[str, Any]]:
    lines = sections.get("experience", [])
    if not lines:
        return []
    blocks = _collect_blocks(lines)
    exps: List[Dict[str, Any]] = []
    for b in blocks:
        block_text = " ".join(b)
        m = DATE_SPAN_RE.search(block_text)
        start = end = None
        if m:
            start = m.group("start")
            end = m.group("end")

        title = company = loc = None
        if b:
            hdr = b[0]
            parts = re.split(r"\s*[|\-–—,]\s*", hdr)
            if len(parts) >= 1:
                title = parts[0].strip() or None
            if len(parts) >= 2:
                company = parts[1].strip() or None
            if len(parts) >= 3:
                loc = parts[2].strip() or None

            if title and DATE_SPAN_RE.search(title) and len(b) > 1:
                hdr2 = b[1]
                parts2 = re.split(r"\s*[|\-–—,]\s*", hdr2)
                title = parts2[0].strip() if len(parts2) > 0 else None
                company = parts2[1].strip() if len(parts2) > 1 else company
                loc = parts2[2].strip() if len(parts2) > 2 else loc

        desc = _parse_bullets(b[1:])
        exps.append({
            "title": title,
            "company": company,
            "location": loc,
            "start_date": start,
            "end_date": end,
            "duration": None,
            "description": desc
        })
    return exps

def _extract_education(sections: Dict[str, List[str]]) -> List[Dict[str, Any]]:
    lines = sections.get("education", [])
    if not lines:
        return []
    blocks = _collect_blocks(lines)
    edus: List[Dict[str, Any]] = []
    for b in blocks:
        degree = institution = loc = grad_year = field = None
        start = end = None

        if b:
            head = b[0].strip()
            mhead = re.match(
                r"^(?P<inst>.+?)\s+(?P<start>(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|\d{1,2}/\d{4}|\d{4})\s*(?:[-–—]|to)\s*(?P<end>(?:Present|Now|Current|present|now|current|\d{4}|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|\d{1,2}/\d{4}))$",
                head
            )
            if mhead:
                institution = mhead.group("inst").strip()
                start = mhead.group("start").strip()
                end = mhead.group("end").strip()
                if len(b) > 1:
                    line2 = b[1].strip()
                    mloc = re.search(r"(.+?)\s+([A-Za-z][A-Za-z .'-]+,\s*[A-Za-z .'-]+)$", line2)
                    if mloc:
                        degree = mloc.group(1).strip()
                        loc = mloc.group(2).strip()
                    else:
                        degree = line2

        if not institution and b:
            hdr = b[0]
            parts = re.split(r"\s*[|\-–—,]\s*", hdr)
            if len(parts) >= 1:
                if re.search(r"(Bachelor|Master|B\.|M\.|PhD|Diploma|Certificate|Postgraduate)", parts[0], re.I):
                    degree = parts[0].strip()
                else:
                    institution = parts[0].strip()
            if len(parts) >= 2:
                if not institution and re.search(r"(University|College|Institute|School)", parts[1], re.I):
                    institution = parts[1].strip()
                else:
                    loc = parts[1].strip()
            if len(parts) >= 3 and not loc:
                loc = parts[2].strip()

        joined = " ".join(b)
        y = re.findall(r"\b(19|20)\d{2}\b", joined)
        if y:
            last = y[-1]
            grad_year = last if isinstance(last, str) else "".join(last)

        edus.append({
            "degree": degree,
            "institution": institution,
            "graduation_year": grad_year,
            "gpa": None,
            "field_of_study": field,
            "location": loc,
            "start_date": start,
            "end_date": end
        })
    return edus

def _extract_projects(sections: Dict[str, List[str]]) -> List[Dict[str, Any]]:
    lines = sections.get("projects", [])
    if not lines:
        return []
    blocks = _collect_blocks(lines)
    res: List[Dict[str, Any]] = []
    for b in blocks:
        name = b[0].strip() if b else None
        desc = " ".join(_parse_bullets(b[1:])) if len(b) > 1 else None
        url = None
        m = URL_RE.search(" ".join(b))
        if m:
            url = _clean_url(m.group(0))
        techs = []
        m2 = re.search(r"(Developed with|Tech|Technolog(?:y|ies)|Stack)\s*[:\-]\s*(.+)", " ".join(b), re.I)
        if m2:
            techs = [t.strip() for t in re.split(r"[,/|]", m2.group(2)) if t.strip()]
        res.append({
            "name": name,
            "description": desc,
            "technologies": techs,
            "url": url
        })
    return res

def _extract_certifications(sections: Dict[str, List[str]]) -> List[Dict[str, Any]]:
    lines = sections.get("certifications", [])
    if not lines:
        return []
    res: List[Dict[str, Any]] = []
    for ln in lines:
        s = ln.strip()
        if not s:
            continue
        parts = re.split(r"\s*[|\-–—]\s*", s)
        name = parts[0].strip() if parts else s
        issuer = parts[1].strip() if len(parts) > 1 else None
        date = parts[2].strip() if len(parts) > 2 else None
        res.append({"name": name, "issuer": issuer, "date": date, "expiry": None})
    return res

def _extract_languages(sections: Dict[str, List[str]]) -> List[str]:
    lines = sections.get("languages", [])
    out: List[str] = []
    for ln in lines:
        out += [x.strip() for x in re.split(r"[,/|;]", ln) if x.strip()]
    return out

def _extract_awards(sections: Dict[str, List[str]]) -> List[str]:
    lines = sections.get("awards", [])
    return [ln.strip() for ln in lines if ln.strip()]

def _extract_summary(sections: Dict[str, List[str]]) -> Optional[str]:
    lines = sections.get("summary", [])
    if not lines:
        return None
    chunk = []
    for ln in lines:
        if ln.strip():
            chunk.append(ln.strip())
        if len(chunk) >= 6:
            break
    return " ".join(chunk) if chunk else None

def _extract_skills_generic(sections: Dict[str, List[str]]) -> Dict[str, List[str]]:
    # Look for skills in multiple section types
    skill_sections = ["skills", "competencies", "core competencies", "key skills", "technical skills", 
                     "professional skills", "expertise", "areas of expertise", "qualifications",
                     "strengths", "core strengths", "abilities", "proficiencies"]
    
    lines = []
    for section_name in skill_sections:
        if section_name in sections:
            lines.extend(sections[section_name])
    
    # Also check if skills are embedded in summary or other sections
    if not lines:
        summary_lines = sections.get("summary", []) + sections.get("objective", []) + sections.get("profile", [])
        for line in summary_lines:
            if any(keyword in line.lower() for keyword in ["skilled in", "proficient in", "experienced in", "expertise in", "knowledge of"]):
                lines.append(line)
    
    buckets = {"technical": [], "soft": [], "domain": []}
    if not lines:
        return buckets

    # Expanded categories for all professions
    TECH_LABELS = {
        # IT & Software
        "technical", "technical skills", "programming", "frontend", "back-end", "backend", "full-stack",
        "cloud", "cloud/devops", "devops", "tools", "other tools", "frameworks", "libraries",
        "databases", "database", "data", "ai/ml & data", "ai/ml", "ai", "ml", "analytics",
        "blockchain", "security", "testing", "qa", "platforms", "infrastructure", "systems",
        "etl", "bi", "visualization", "mobile", "ios", "android", "automation",
        # Engineering & Manufacturing
        "engineering", "cad", "design software", "modeling", "simulation", "autocad", "solidworks",
        "matlab", "plc", "scada", "cnc", "lean manufacturing", "six sigma", "quality control",
        # Healthcare & Medical
        "medical software", "emr", "ehr", "epic", "cerner", "medical devices", "laboratory equipment",
        "diagnostic tools", "imaging", "radiology", "pathology", "clinical research",
        # Finance & Accounting
        "financial software", "excel", "quickbooks", "sap", "oracle", "bloomberg", "financial modeling",
        "risk management", "trading platforms", "accounting software", "erp", "crm",
        # Creative & Design
        "adobe", "photoshop", "illustrator", "indesign", "figma", "sketch", "3d modeling",
        "video editing", "audio editing", "graphic design", "web design", "ux/ui",
        # Scientific & Research
        "laboratory techniques", "research methods", "statistical analysis", "spss", "r", "python",
        "microscopy", "spectroscopy", "chromatography", "pcr", "cell culture"
    }
    
    SOFT_LABELS = {
        "soft", "soft skills", "communication", "leadership", "interpersonal", "collaboration",
        "teamwork", "problem solving", "time management", "organization", "stakeholder management",
        "customer service", "sales", "negotiation", "presentation", "public speaking",
        "conflict resolution", "mentoring", "coaching", "training", "facilitation",
        "critical thinking", "analytical thinking", "creativity", "innovation", "adaptability",
        "flexibility", "resilience", "emotional intelligence", "cultural awareness",
        "project management", "team building", "decision making", "strategic thinking"
    }
    
    DOMAIN_LABELS = {
        "domain", "business", "industry", "methodologies", "practices", "certifications",
        "compliance", "regulatory", "finance", "marketing", "sales", "operations", "supply chain",
        "hr", "human resources", "healthcare", "education", "real estate", "banking", "hospitality", "retail",
        # Healthcare specialties
        "cardiology", "oncology", "pediatrics", "surgery", "nursing", "pharmacy", "physical therapy",
        "occupational therapy", "radiology", "pathology", "anesthesiology", "emergency medicine",
        # Legal
        "litigation", "corporate law", "family law", "criminal law", "intellectual property",
        "contract law", "real estate law", "tax law", "employment law", "immigration law",
        # Education
        "curriculum development", "lesson planning", "classroom management", "assessment",
        "special education", "esl", "stem", "early childhood", "higher education",
        # Business & Management
        "strategic planning", "business development", "market research", "brand management",
        "digital marketing", "content marketing", "social media", "seo", "sem", "ppc",
        "supply chain management", "logistics", "procurement", "vendor management",
        # Manufacturing & Engineering
        "process improvement", "quality assurance", "regulatory compliance", "safety management",
        "environmental compliance", "iso standards", "fda regulations", "gmp", "haccp",
        # Creative Industries
        "brand identity", "visual communication", "user experience", "user interface",
        "motion graphics", "typography", "color theory", "composition", "storytelling"
    }

    def label_to_bucket(label: str) -> str:
        lab = label.strip().lower().rstrip(":")
        if lab in TECH_LABELS:
            return "technical"
        if lab in SOFT_LABELS:
            return "soft"
        if lab in DOMAIN_LABELS:
            return "domain"
        if re.search(r"(programming|frontend|backend|cloud|devops|tools|framework|library|database|data|ai|ml|blockchain|security|testing|qa|platform|infra|etl|bi|viz|mobile|automation)", lab):
            return "technical"
        if re.search(r"(communication|leadership|team|stakeholder|management|problem|time|organization|customer|sales|negotiation|presentation)", lab):
            return "soft"
        return "domain"

    current = "technical"  # default inside a Skills section

    for raw in lines:
        ln = raw.strip()
        if not ln:
            continue

        # skip generic headings
        low = ln.lower().rstrip(":")
        if low in ("skills", "technical skills", "soft skills", "core skills", "competencies", "strengths",
                   "technical competencies", "areas of expertise", "key skills", "core competencies"):
            continue

        # If this is a label line (ends with ':'), switch bucket and continue
        if ln.endswith(":"):
            current = label_to_bucket(ln)
            continue

        # Split items (comma/pipe/slash) but also handle single items per line
        parts = [x.strip() for x in re.split(r"[,/|;]", ln) if x.strip()]
        if not parts:
            continue

        buckets[current].extend(parts)

    # Deduplicate while preserving original order
    for k in buckets:
        seen = set()
        uniq = []
        for item in buckets[k]:
            key = item.lower()
            if key not in seen:
                seen.add(key)
                uniq.append(item)
        buckets[k] = uniq

    # Enhanced fallback: try to extract skills from free text
    if not any(buckets.values()):
        flat = []
        for ln in lines:
            # Handle bullet points and various separators
            clean_line = re.sub(r'^[•◦▪■–—-]\s*', '', ln.strip())
            if clean_line:
                # Split on common separators but also handle single skills per line
                parts = [x.strip() for x in re.split(r'[,/|;•]', clean_line) if x.strip()]
                if not parts and clean_line:  # Single skill on line
                    parts = [clean_line]
                flat.extend(parts)
        
        # Clean and categorize extracted skills
        seen = set()
        for skill in flat:
            skill_clean = skill.strip()
            skill_lower = skill_clean.lower()
            
            # Skip generic headers and empty strings
            if (skill_lower in {"skills", "technical skills", "soft skills", "core skills", 
                               "competencies", "key skills", "areas of expertise"} or 
                len(skill_clean) < 2 or skill_lower in seen):
                continue
                
            seen.add(skill_lower)
            
            # Smart categorization based on skill content
            if any(tech_word in skill_lower for tech_word in [
                'software', 'programming', 'coding', 'development', 'database', 'system',
                'technology', 'technical', 'computer', 'digital', 'web', 'mobile', 'app',
                'cloud', 'server', 'network', 'security', 'automation', 'analytics',
                'microsoft', 'adobe', 'google', 'amazon', 'oracle', 'sap', 'salesforce'
            ]):
                buckets["technical"].append(skill_clean)
            elif any(soft_word in skill_lower for soft_word in [
                'communication', 'leadership', 'team', 'management', 'organization',
                'problem', 'critical', 'analytical', 'creative', 'interpersonal',
                'presentation', 'negotiation', 'customer', 'client', 'service'
            ]):
                buckets["soft"].append(skill_clean)
            else:
                buckets["domain"].append(skill_clean)

    return buckets

# ----------------------------
# File text extraction
# ----------------------------

def _extract_text_from_pdf_bytes(file_content: bytes) -> str:
    """
    Extract text and hyperlink URIs from a PDF.
    We keep layout via pdfplumber when possible and also harvest link annotations
    (LinkedIn/GitHub/Website), appending them so URL regexes can capture them.
    """
    import io
    link_uris = set()
    text = ""

    # Try pdfplumber first (can expose page.hyperlinks)
    try:
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            pages = []
            for pg in pdf.pages:
                # page text
                t = pg.extract_text(x_tolerance=1.5, y_tolerance=2.0)
                if t:
                    pages.append(t)

                # hyperlink annotations (if available)
                try:
                    for h in getattr(pg, "hyperlinks", []) or []:
                        uri = h.get("uri") or h.get("target")
                        if uri and isinstance(uri, str):
                            link_uris.add(uri.strip())
                except Exception:
                    pass

            if pages:
                text = "\n\n".join(pages)
    except Exception as e:
        print(f"pdfplumber failed: {e}")

    # Fallback: PyPDF2 for text + /Annots URIs
    if not text:
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            out = []
            for p in reader.pages:
                try:
                    tt = p.extract_text() or ""
                except Exception:
                    tt = ""
                if tt:
                    out.append(tt)

                # Extract URI from annotations
                try:
                    annots = p.get("/Annots") or p.get("/Annots".encode())
                    if annots:
                        for a in annots:
                            obj = a.get_object()
                            adict = obj.get("/A") if obj else None
                            uri = adict.get("/URI") if adict else None
                            if uri and isinstance(uri, str):
                                link_uris.add(uri.strip())
                except Exception:
                    pass

            text = "\n\n".join(out)
        except Exception as e:
            print(f"PyPDF2 failed: {e}")
            text = ""

    # Append discovered links so downstream regex can see them
    if link_uris:
        cleaned = [_clean_url(u) for u in sorted(link_uris)]
        text = (text or "").rstrip() + "\n\n" + "\n".join(f"LINK: {u}" for u in cleaned)

    return text

def _extract_text_from_docx_bytes(file_content: bytes) -> str:
    try:
        import io
        d = docx.Document(io.BytesIO(file_content))
        parts = []
        for p in d.paragraphs:
            parts.append(p.text)
        # include simple tables
        for tbl in d.tables:
            for row in tbl.rows:
                cells = [c.text for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as e:
        print(f"DOCX extraction failed: {e}")
        return ""

# ----------------------------
# Main Parser Class
# ----------------------------

class EnhancedResumeParser:
    def __init__(self):
        self.langchain_available = LANGCHAIN_AVAILABLE

        if self.langchain_available:
            self.openai_api_key = os.getenv("OPENAI_API_KEY")
            if not self.openai_api_key:
                print("Warning: OPENAI_API_KEY not found, falling back to basic parsing")
                self.langchain_available = False
            else:
                try:
                    # Keep your original ChatOpenAI usage
                    self.llm = ChatOpenAI(
                        model=os.getenv("RESUME_PARSER_MODEL", "gpt-3.5-turbo"),
                        temperature=0,
                        openai_api_key=self.openai_api_key
                    )
                    # Keep PydanticOutputParser
                    self.parser = PydanticOutputParser(pydantic_object=ParsedResumeStructure)
                    # Stronger, but same PromptTemplate interface
                    self.prompt_template = PromptTemplate(
                        template=(
                            "You are an expert resume parser. Extract structured information from the resume text below.\n"
                            "Return ONLY fields described in the schema via the format instructions.\n"
                            "Guidelines:\n"
                            "- Preserve bullets as array items for descriptions.\n"
                            "- Keep date strings as they appear (e.g., 'Jan 2024 – Present', '03/2020 – 06/2021').\n"
                            "- Extract email/phone/LinkedIn/GitHub/website from ANYWHERE in the document (not just header).\n"
                            "- Education may appear as two lines (Institution + dates, then Degree + location). Parse accordingly.\n"
                            "- If information is missing, use nulls or empty arrays. Do not invent facts.\n\n"
                            "Resume Text:\n{resume_text}\n\n{format_instructions}\n\nParsed Resume:\n"
                        ),
                        input_variables=["resume_text"],
                        partial_variables={"format_instructions": self.parser.get_format_instructions()}
                    )
                except Exception as e:
                    print(f"Error initializing LangChain: {e}, falling back to basic parsing")
                    self.langchain_available = False

    async def extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file (preserving layout where possible)"""
        return _normalize_text_keep_layout(_extract_text_from_pdf_bytes(file_content))

    async def extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file (preserving layout)"""
        return _normalize_text_keep_layout(_extract_text_from_docx_bytes(file_content))

    async def parse_resume(self, file_content: bytes, content_type: str) -> Dict[str, Any]:
        """Parse resume and return structured data (same output shape)"""
        try:
            # Extract text based on file type
            if content_type == "application/pdf":
                text = await self.extract_text_from_pdf(file_content)
            elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                text = await self.extract_text_from_docx(file_content)
            else:
                raise ValueError(f"Unsupported file type: {content_type}")

            if not text.strip():
                raise ValueError("Could not extract text from the resume")

            cleaned_text = text  # layout preserved

            # Try enhanced parsing if LangChain is available
            if self.langchain_available:
                try:
                    prompt = self.prompt_template.format(resume_text=cleaned_text)
                    response = self.llm.predict(prompt)  # legacy-compatible; you can swap to .invoke() later
                    parsed_resume = self.parser.parse(response)
                    result = parsed_resume.dict()

                    # Ensure keys exist even if model left them empty
                    result.setdefault("skills", {"technical": [], "soft": [], "domain": []})
                    result.setdefault("experience", [])
                    result.setdefault("education", [])
                    result.setdefault("projects", [])
                    result.setdefault("certifications", [])
                    result.setdefault("languages", [])
                    result.setdefault("awards", [])

                    # Merge LLM output with robust rule-based signals (backfill links/edu/exp + location)
                    result = self._merge_with_rules(result, cleaned_text)

                    # Enrich with metadata and stats
                    result["raw_text"] = cleaned_text
                    result["parsing_method"] = "openai_langchain"
                    result["parsed_at"] = datetime.utcnow().isoformat()
                    result["statistics"] = {
                        "total_experience_years": self._calculate_total_experience(result.get("experience", [])),
                        "skill_count": sum(len(v) for v in (result.get("skills") or {}).values()) if isinstance(result.get("skills"), dict) else 0,
                        "education_count": len(result.get("education", [])),
                        "project_count": len(result.get("projects", [])),
                        "certification_count": len(result.get("certifications", []))
                    }
                    return result
                except Exception as e:
                    print(f"Enhanced parsing failed: {e}, falling back to rule-based parsing")

            # Rule-based parsing fallback (strong)
            return await self._basic_parse_fallback_with_sections(cleaned_text)

        except Exception as e:
            print(f"Resume parsing failed: {e}")
            # Final fallback
            return await self._basic_parse_fallback_with_sections("")

    def _merge_with_rules(self, result: Dict[str, Any], cleaned_text: str) -> Dict[str, Any]:
        """Backfill/repair: links anywhere, skills, 2-line education, misaligned experience, top-level location."""
        try:
            # Parse sections for rule-based extraction
            sections = _split_sections(cleaned_text)
            
            # Extract personal info if missing or incomplete
            if not result.get("personal_info") or not any(result["personal_info"].values()):
                result["personal_info"] = _extract_personal_info(sections)
            else:
                # Backfill missing fields
                rule_info = _extract_personal_info(sections)
                for key, value in rule_info.items():
                    if value and not result["personal_info"].get(key):
                        result["personal_info"][key] = value
            
            # Extract skills if missing or empty - enhanced logic
            skills_empty = (
                not result.get("skills") or 
                (isinstance(result["skills"], dict) and not any(result["skills"].values())) or
                (isinstance(result["skills"], list) and not result["skills"])
            )
            
            if skills_empty:
                extracted_skills = _extract_skills_generic(sections)
                result["skills"] = extracted_skills
                print(f"Extracted skills: {extracted_skills}")  # Debug logging
            
            # Convert old list format to new dict format if needed
            if isinstance(result.get("skills"), list):
                old_skills = result["skills"]
                result["skills"] = {"technical": old_skills, "soft": [], "domain": []}
            
            # Extract experience if missing
            if not result.get("experience"):
                result["experience"] = _extract_experience(sections)
            
            # Extract education if missing
            if not result.get("education"):
                result["education"] = _extract_education(sections)
            
            # Extract other sections if missing
            if not result.get("projects"):
                result["projects"] = _extract_projects(sections)
            
            if not result.get("certifications"):
                result["certifications"] = _extract_certifications(sections)
            
            if not result.get("languages"):
                result["languages"] = _extract_languages(sections)
            
            if not result.get("awards"):
                result["awards"] = _extract_awards(sections)
            
            if not result.get("summary"):
                result["summary"] = _extract_summary(sections)
            
            return result
        except Exception as e:
            print(f"Error in _merge_with_rules: {e}")
            return result

    def _clean_text(self, text: str) -> str:
        """(Kept for compatibility; unused to avoid destroying structure)"""
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[^\w\s@.\-:/]', ' ', text)
        return text.strip()

    def _calculate_total_experience(self, experiences: List[Dict]) -> float:
        """Calculate total years of experience using parsed date spans when available"""
        total_months = 0
        for exp in experiences:
            try:
                s = (exp or {}).get("start_date") or ""
                e = (exp or {}).get("end_date") or ""
                start_dt = _parse_month_year(s)
                if e and re.search(r"(present|current|now)", e, re.I):
                    end_dt = None
                else:
                    end_dt = _parse_month_year(e)
                total_months += _months_between(start_dt, end_dt)
            except Exception:
                continue
        return round(total_months / 12.0, 1)

    async def _basic_parse_fallback(self, file_content: bytes, content_type: str) -> Dict[str, Any]:
        """Legacy fallback (kept for compatibility)"""
        try:
            if content_type == "application/pdf":
                text = await self.extract_text_from_pdf(file_content)
            else:
                text = await self.extract_text_from_docx(file_content)

            result = {
                "personal_info": self._extract_basic_personal_info(text),
                "skills": {"technical": self._extract_basic_skills(text)},
                "raw_text": text,
                "parsing_method": "basic_fallback",
                "parsed_at": datetime.utcnow().isoformat(),
                "experience": [],
                "education": [],
                "projects": [],
                "certifications": [],
                "languages": [],
                "awards": []
            }
            return result
        except Exception as e:
            print(f"Basic parsing also failed: {e}")
            return {
                "error": str(e),
                "parsing_method": "failed",
                "parsed_at": datetime.utcnow().isoformat()
            }

    async def _basic_parse_fallback_with_sections(self, text: str) -> Dict[str, Any]:
        """Strong rule-based fallback using section parsing; same output shape."""
        try:
            sections = _split_sections(text or "")

            personal_info = _extract_personal_info(sections)
            summary = _extract_summary(sections)
            skills = _extract_skills_generic(sections)
            experience = _extract_experience(sections)
            education = _extract_education(sections)
            projects = _extract_projects(sections)
            certifications = _extract_certifications(sections)
            languages = _extract_languages(sections)
            awards = _extract_awards(sections)

            # Backfill top-level location from available content
            if not personal_info.get("location"):
                personal_info["location"] = _majority_vote_location(sections.get("header", []), text, experience, education)

            result = {
                "personal_info": personal_info,
                "summary": summary,
                "skills": skills,
                "experience": experience,
                "education": education,
                "projects": projects,
                "certifications": certifications,
                "languages": languages,
                "awards": awards,
                "raw_text": text,
                "parsing_method": "rules_only",
                "parsed_at": datetime.utcnow().isoformat(),
                "statistics": {
                    "total_experience_years": self._calculate_total_experience(experience),
                    "skill_count": sum(len(v) for v in (skills or {}).values()) if isinstance(skills, dict) else 0,
                    "education_count": len(education),
                    "project_count": len(projects),
                    "certification_count": len(certifications)
                }
            }
            return result
        except Exception as e:
            print(f"Rule-based parsing failed: {e}")
            return {
                "error": str(e),
                "parsing_method": "failed",
                "parsed_at": datetime.utcnow().isoformat()
            }

    def _extract_basic_personal_info(self, text: str) -> Dict[str, Any]:
        """Basic personal info extraction (legacy; used only by old fallback)"""
        info = {}
        email_match = EMAIL_RE.search(text)
        if email_match:
            info["email"] = email_match.group()
        phone_match = PHONE_RE.search(text)
        if phone_match:
            info["phone"] = phone_match.group()
        return info

    def _extract_basic_skills(self, text: str) -> Dict[str, List[str]]:
        """Enhanced basic skills extraction for all professions"""
        # Comprehensive skill database for all professions
        all_skills = {
            'technical': [
                # Programming & Software
                'python', 'javascript', 'java', 'react', 'node.js', 'sql', 'html', 'css',
                'aws', 'docker', 'kubernetes', 'git', 'mongodb', 'postgresql', 'redis',
                'angular', 'vue.js', 'typescript', 'c++', 'c#', '.net', 'php', 'ruby',
                # Design & Creative
                'photoshop', 'illustrator', 'indesign', 'figma', 'sketch', 'autocad',
                'solidworks', 'maya', 'blender', 'after effects', 'premiere pro',
                # Business Software
                'excel', 'powerpoint', 'word', 'outlook', 'sharepoint', 'teams',
                'salesforce', 'hubspot', 'quickbooks', 'sap', 'oracle', 'tableau',
                # Medical & Healthcare
                'epic', 'cerner', 'meditech', 'allscripts', 'emr', 'ehr', 'pacs',
                # Engineering
                'matlab', 'labview', 'plc', 'scada', 'cnc', 'cad', 'fem analysis'
            ],
            'soft': [
                'communication', 'leadership', 'teamwork', 'problem solving',
                'time management', 'organization', 'critical thinking', 'creativity',
                'adaptability', 'collaboration', 'presentation', 'negotiation',
                'customer service', 'project management', 'analytical thinking',
                'decision making', 'conflict resolution', 'mentoring', 'coaching'
            ],
            'domain': [
                # Healthcare
                'patient care', 'clinical research', 'medical coding', 'hipaa',
                'pharmacology', 'anatomy', 'physiology', 'pathology',
                # Finance
                'financial analysis', 'risk management', 'compliance', 'audit',
                'investment', 'portfolio management', 'derivatives', 'forex',
                # Legal
                'litigation', 'contract law', 'intellectual property', 'compliance',
                # Education
                'curriculum development', 'lesson planning', 'assessment',
                'classroom management', 'special education', 'esl',
                # Marketing
                'digital marketing', 'seo', 'sem', 'social media', 'content marketing',
                'brand management', 'market research', 'analytics'
            ]
        }
        
        found_skills = {'technical': [], 'soft': [], 'domain': []}
        text_lower = text.lower()
        
        for category, skills_list in all_skills.items():
            for skill in skills_list:
                if skill.lower() in text_lower:
                    found_skills[category].append(skill)
        
        return found_skills

# Initialize enhanced resume parser with error handling
try:
    enhanced_resume_parser = EnhancedResumeParser()
except Exception as e:
    print(f"Warning: Could not initialize enhanced resume parser: {e}")
    print("Falling back to basic parsing only")
    enhanced_resume_parser = None
