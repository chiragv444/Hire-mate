import PyPDF2
import docx
import os
import uuid
import re
# import spacy  # Temporarily disabled due to dependency conflicts
import nltk
from typing import Dict, List, Any, Optional
from datetime import datetime
import aiofiles
from pathlib import Path
import pdfplumber
import fitz  # PyMuPDF
from textblob import TextBlob
from fuzzywuzzy import fuzz
from .enhanced_resume_parser import enhanced_resume_parser

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')

class ResumeParser:
    def __init__(self):
        # Load spaCy model (temporarily disabled)
        # try:
        #     self.nlp = spacy.load("en_core_web_sm")
        # except OSError:
        #     # Download if not available
        #     os.system("python -m spacy download en_core_web_sm")
        #     self.nlp = spacy.load("en_core_web_sm")
        self.nlp = None  # Temporarily disabled
        
        # Comprehensive skills database
        self.skills_keywords = {
            'programming': [
                'python', 'javascript', 'java', 'c++', 'c#', 'php', 'ruby', 'go', 'rust', 'swift', 'kotlin',
                'typescript', 'scala', 'r', 'matlab', 'perl', 'shell', 'bash', 'powershell', 'vba',
                'html', 'css', 'sass', 'less', 'xml', 'json', 'yaml'
            ],
            'frameworks': [
                'react', 'angular', 'vue', 'django', 'flask', 'spring', 'express', 'laravel', 'rails',
                'nextjs', 'nuxtjs', 'gatsby', 'svelte', 'ember', 'backbone', 'jquery', 'bootstrap',
                'tailwind', 'material-ui', 'ant design', 'chakra ui', 'fastapi', 'tornado', 'pyramid'
            ],
            'databases': [
                'mysql', 'postgresql', 'mongodb', 'redis', 'sqlite', 'oracle', 'sql server', 'mariadb',
                'cassandra', 'dynamodb', 'elasticsearch', 'neo4j', 'couchdb', 'firebase', 'supabase'
            ],
            'cloud': [
                'aws', 'azure', 'gcp', 'google cloud', 'docker', 'kubernetes', 'terraform', 'jenkins',
                'gitlab ci', 'github actions', 'circleci', 'travis ci', 'ansible', 'puppet', 'chef',
                'vagrant', 'helm', 'istio', 'prometheus', 'grafana', 'elk stack', 'datadog', 'newrelic'
            ],
            'tools': [
                'git', 'github', 'gitlab', 'bitbucket', 'jira', 'confluence', 'slack', 'trello', 'figma',
                'sketch', 'adobe xd', 'photoshop', 'illustrator', 'postman', 'insomnia', 'swagger',
                'vs code', 'intellij', 'eclipse', 'vim', 'emacs', 'sublime text', 'atom'
            ],
            'data_science': [
                'pandas', 'numpy', 'scipy', 'scikit-learn', 'tensorflow', 'pytorch', 'keras',
                'matplotlib', 'seaborn', 'plotly', 'jupyter', 'anaconda', 'spark', 'hadoop',
                'tableau', 'power bi', 'qlik', 'looker', 'dbt', 'airflow', 'kafka', 'storm'
            ],
            'mobile': [
                'ios', 'android', 'react native', 'flutter', 'xamarin', 'ionic', 'cordova',
                'swift', 'objective-c', 'kotlin', 'java', 'dart'
            ],
            'languages': ['english', 'spanish', 'french', 'german', 'chinese', 'japanese', 'korean', 'portuguese', 'italian', 'russian', 'arabic', 'hindi'],
            'soft_skills': [
                'leadership', 'communication', 'teamwork', 'problem solving', 'analytical', 'creative',
                'project management', 'agile', 'scrum', 'kanban', 'time management', 'critical thinking',
                'adaptability', 'collaboration', 'mentoring', 'public speaking', 'negotiation'
            ],
            'certifications': [
                'aws certified', 'azure certified', 'gcp certified', 'pmp', 'scrum master', 'cissp',
                'comptia', 'cisco', 'microsoft certified', 'oracle certified', 'salesforce certified'
            ]
        }
        
        # Experience patterns
        self.experience_patterns = [
            r'(\d{4})\s*[-–]\s*(\d{4}|present|current)',
            r'(\d{1,2}/\d{4})\s*[-–]\s*(\d{1,2}/\d{4}|present|current)',
            r'(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\s+\d{4}\s*[-–]\s*(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\s+\d{4}',
        ]
    
    async def save_uploaded_file(self, file_content: bytes, original_filename: str, upload_dir: str) -> Dict[str, Any]:
        """Save uploaded file and return metadata"""
        # Create unique filename
        file_extension = Path(original_filename).suffix.lower()
        unique_filename = f"{uuid.uuid4()}{file_extension}"
        
        # Ensure upload directory exists
        os.makedirs(upload_dir, exist_ok=True)
        file_path = os.path.join(upload_dir, unique_filename)
        
        # Save file
        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(file_content)
        
        return {
            'filename': unique_filename,
            'original_name': original_filename,
            'file_size': len(file_content),
            'file_type': file_extension,
            'file_path': file_path,
            'upload_date': datetime.now()
        }
    
    async def parse_resume(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """Parse resume file and extract structured data using enhanced parser"""
        try:
            # Read file content
            async with aiofiles.open(file_path, 'rb') as f:
                file_content = await f.read()
            
            # Use enhanced resume parser if available
            if enhanced_resume_parser:
                try:
                    parsed_data = enhanced_resume_parser.parse_resume(file_content, file_type)
                    
                    # Ensure skills are in the correct format
                    if isinstance(parsed_data.get('skills'), list):
                        # Convert old list format to new dict format
                        old_skills = parsed_data['skills']
                        parsed_data['skills'] = {
                            'technical': old_skills,
                            'soft': [],
                            'domain': []
                        }
                    elif not isinstance(parsed_data.get('skills'), dict):
                        # Fallback to empty skills structure
                        parsed_data['skills'] = {
                            'technical': [],
                            'soft': [],
                            'domain': []
                        }
                    
                    return parsed_data
                    
                except Exception as e:
                    print(f"Enhanced parser failed, falling back to basic parser: {e}")
                    # Fall through to basic parser
            
            # Fallback to basic parsing
            raw_text = await self._extract_text(file_path, file_type)
            cleaned_text = self._clean_text(raw_text)
            
            # Use enhanced basic skills extraction
            basic_skills = self._extract_skills_enhanced(cleaned_text)
            
            parsed_data = {
                'raw_text': cleaned_text,
                'skills': basic_skills,
                'experience': self._extract_experience(cleaned_text),
                'education': self._extract_education(cleaned_text),
                'contact_info': self._extract_contact_info(cleaned_text),
                'summary': self._extract_summary(cleaned_text),
                'languages': self._extract_languages(cleaned_text),
                'certifications': self._extract_certifications(cleaned_text)
            }
            
            return parsed_data
            
        except Exception as e:
            raise Exception(f"Error parsing resume: {str(e)}")
    
    async def parse_resume_from_storage(self, storage_path: str) -> Dict[str, Any]:
        """Parse resume file from Firebase Storage path"""
        try:
            # For now, assume the file is already downloaded locally
            # In a full implementation, you'd download from Firebase Storage first
            # This is a simplified version that works with local file paths
            
            # Determine file type from path
            file_extension = os.path.splitext(storage_path)[1].lower()
            file_type = 'pdf' if file_extension == '.pdf' else 'docx' if file_extension in ['.docx', '.doc'] else 'txt'
            
            # For Firebase Storage, you would typically:
            # 1. Download the file from storage_path
            # 2. Save it temporarily
            # 3. Parse it
            # 4. Clean up the temporary file
            
            # For now, we'll assume the storage_path is a local file path
            # This will need to be updated when Firebase Storage integration is complete
            return await self.parse_resume(storage_path, file_type)
            
        except Exception as e:
            print(f"Error parsing resume from storage: {e}")
            raise
    
    async def _extract_text(self, file_path: str, file_type: str) -> str:
        """Extract text from PDF or DOCX file"""
        if file_type == '.pdf':
            return await self._extract_pdf_text(file_path)
        elif file_type == '.docx':
            return await self._extract_docx_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    async def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file using multiple methods for better accuracy"""
        text = ""
        
        # Method 1: Try pdfplumber first (better for complex layouts)
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            if text.strip():
                return text
        except Exception as e:
            print(f"pdfplumber failed: {e}")
        
        # Method 2: Try PyMuPDF as fallback
        try:
            doc = fitz.open(file_path)
            for page_num in range(doc.page_count):
                page = doc[page_num]
                text += page.get_text() + "\n"
            doc.close()
            
            if text.strip():
                return text
        except Exception as e:
            print(f"PyMuPDF failed: {e}")
        
        # Method 3: PyPDF2 as final fallback
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            print(f"PyPDF2 failed: {e}")
        
        return text if text.strip() else "Could not extract text from PDF"
    
    async def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text
        except Exception as e:
            raise Exception(f"Error extracting DOCX text: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters but keep important ones
        text = re.sub(r'[^\w\s\-\.\,\;\:\!\?\(\)\[\]\{\}]', ' ', text)
        # Normalize line breaks
        text = re.sub(r'\n+', '\n', text)
        return text.strip()
    
    def _extract_skills_enhanced(self, text: str) -> Dict[str, List[str]]:
        """Enhanced skills extraction for all professions"""
        skills = {'technical': [], 'soft': [], 'domain': []}
        text_lower = text.lower()
        
        # Comprehensive skill categories for all professions
        skill_categories = {
            'technical': {
                # Programming & Software
                'python', 'javascript', 'java', 'react', 'node.js', 'sql', 'html', 'css',
                'aws', 'docker', 'kubernetes', 'git', 'mongodb', 'postgresql', 'redis',
                'angular', 'vue.js', 'typescript', 'c++', 'c#', '.net', 'php', 'ruby',
                # Design & Creative Software
                'photoshop', 'illustrator', 'indesign', 'figma', 'sketch', 'autocad',
                'solidworks', 'maya', 'blender', 'after effects', 'premiere pro',
                # Business & Office Software
                'excel', 'powerpoint', 'word', 'outlook', 'sharepoint', 'teams',
                'salesforce', 'hubspot', 'quickbooks', 'sap', 'oracle', 'tableau',
                # Medical & Healthcare Software
                'epic', 'cerner', 'meditech', 'allscripts', 'emr', 'ehr', 'pacs',
                # Engineering & Scientific
                'matlab', 'labview', 'plc', 'scada', 'cnc', 'cad', 'fem analysis',
                'spss', 'r', 'stata', 'minitab', 'origin'
            },
            'soft': {
                'communication', 'leadership', 'teamwork', 'problem solving',
                'time management', 'organization', 'critical thinking', 'creativity',
                'adaptability', 'collaboration', 'presentation', 'negotiation',
                'customer service', 'project management', 'analytical thinking',
                'decision making', 'conflict resolution', 'mentoring', 'coaching',
                'public speaking', 'interpersonal skills', 'emotional intelligence'
            },
            'domain': {
                # Healthcare
                'patient care', 'clinical research', 'medical coding', 'hipaa',
                'pharmacology', 'anatomy', 'physiology', 'pathology', 'nursing',
                'surgery', 'radiology', 'cardiology', 'oncology', 'pediatrics',
                # Finance & Accounting
                'financial analysis', 'risk management', 'compliance', 'audit',
                'investment', 'portfolio management', 'derivatives', 'forex',
                'accounting', 'bookkeeping', 'tax preparation', 'budgeting',
                # Legal
                'litigation', 'contract law', 'intellectual property', 'compliance',
                'corporate law', 'family law', 'criminal law', 'real estate law',
                # Education
                'curriculum development', 'lesson planning', 'assessment',
                'classroom management', 'special education', 'esl', 'stem',
                # Marketing & Sales
                'digital marketing', 'seo', 'sem', 'social media', 'content marketing',
                'brand management', 'market research', 'analytics', 'sales',
                # Engineering & Manufacturing
                'process improvement', 'quality assurance', 'lean manufacturing',
                'six sigma', 'iso standards', 'regulatory compliance', 'safety management'
            }
        }
        
        # Extract skills from each category
        for category, skill_set in skill_categories.items():
            for skill in skill_set:
                if skill.lower() in text_lower:
                    skills[category].append(skill.title())
        
        # Remove duplicates while preserving order
        for category in skills:
            skills[category] = list(dict.fromkeys(skills[category]))
        
        return skills
    
    def _extract_skills(self, text: str) -> Dict[str, List[str]]:
        """Legacy method - now calls enhanced extraction"""
        return self._extract_skills_enhanced(text)
    
    def _extract_experience(self, text: str) -> List[Dict[str, Any]]:
        """Extract work experience from resume"""
        experience = []
        
        # Split text into sections
        sections = re.split(r'\n{2,}', text)
        
        for section in sections:
            # Look for experience patterns
            for pattern in self.experience_patterns:
                matches = re.finditer(pattern, section, re.IGNORECASE)
                for match in matches:
                    # Extract company and role information
                    lines = section.split('\n')
                    for line in lines:
                        if any(keyword in line.lower() for keyword in ['inc', 'corp', 'ltd', 'company', 'llc']):
                            experience.append({
                                'company': line.strip(),
                                'duration': match.group(0),
                                'description': section.strip()
                            })
                            break
        
        return experience
    
    def _extract_education(self, text: str) -> List[Dict[str, Any]]:
        """Extract education information"""
        education = []
        
        # Look for education keywords
        education_keywords = ['education', 'academic', 'degree', 'university', 'college', 'school']
        
        lines = text.split('\n')
        for i, line in enumerate(lines):
            if any(keyword in line.lower() for keyword in education_keywords):
                # Extract education details
                education_info = {
                    'institution': line.strip(),
                    'details': lines[i+1].strip() if i+1 < len(lines) else ''
                }
                education.append(education_info)
        
        return education
    
    def _extract_contact_info(self, text: str) -> Dict[str, str]:
        """Extract contact information"""
        contact_info = {}
        
        # Email pattern
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_match = re.search(email_pattern, text)
        if email_match:
            contact_info['email'] = email_match.group()
        
        # Phone pattern
        phone_pattern = r'(\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        phone_match = re.search(phone_pattern, text)
        if phone_match:
            contact_info['phone'] = phone_match.group()
        
        # LinkedIn pattern
        linkedin_pattern = r'linkedin\.com/in/[\w-]+'
        linkedin_match = re.search(linkedin_pattern, text)
        if linkedin_match:
            contact_info['linkedin'] = linkedin_match.group()
        
        return contact_info
    
    def _extract_summary(self, text: str) -> Optional[str]:
        """Extract summary/objective section"""
        summary_keywords = ['summary', 'objective', 'profile', 'overview']
        
        paragraphs = text.split('\n\n')
        for paragraph in paragraphs:
            if any(keyword in paragraph.lower() for keyword in summary_keywords):
                return paragraph.strip()
        
        return None
    
    def _extract_languages(self, text: str) -> List[str]:
        """Extract languages from resume"""
        languages = []
        text_lower = text.lower()
        
        for language in self.skills_keywords['languages']:
            if language in text_lower:
                languages.append(language.title())
        
        return languages
    
    def _extract_certifications(self, text: str) -> List[str]:
        """Extract certifications from resume"""
        certifications = []
        
        cert_keywords = ['certification', 'certified', 'certificate', 'license']
        
        lines = text.split('\n')
        for line in lines:
            if any(keyword in line.lower() for keyword in cert_keywords):
                certifications.append(line.strip())
        
        return certifications

# Initialize resume parser
resume_parser = ResumeParser() 